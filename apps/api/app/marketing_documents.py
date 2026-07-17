"""Export a saved marketing plan as an editable Word file or portable PDF."""

from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from typing import Literal, cast

from docx import Document
from docx.document import Document as DocxDocument
from docx.shared import Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import (
    Flowable,
    KeepTogether,
    ListFlowable,
    ListItem,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from app.marketing import MarketingPlanRequest, MarketingPlanResponse
from app.platform_exports import PlatformValidationError
from app.schemas import MarketingPlanDetailRead, MarketingPlanVersionRead

DocumentFormat = Literal["docx", "pdf"]


@dataclass(frozen=True, slots=True)
class MarketingPlanDocument:
    filename: str
    media_type: str
    content: bytes


def export_marketing_plan_document(
    plan: MarketingPlanDetailRead,
    output_format: DocumentFormat,
    *,
    version_id: str | None = None,
) -> MarketingPlanDocument:
    version = _select_version(plan, version_id)
    request = MarketingPlanRequest.model_validate(version.request_payload)
    content = MarketingPlanResponse.model_validate(version.content)
    filename = f"heyu-marketing-plan-v{version.version_number}.{output_format}"
    if output_format == "docx":
        payload = _build_docx(plan, version, request, content)
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif output_format == "pdf":
        payload = _build_pdf(plan, version, request, content)
        media_type = "application/pdf"
    else:
        raise PlatformValidationError("document format must be docx or pdf")
    return MarketingPlanDocument(filename=filename, media_type=media_type, content=payload)


def _select_version(
    plan: MarketingPlanDetailRead,
    version_id: str | None,
) -> MarketingPlanVersionRead:
    if version_id is None:
        return plan.current_version
    selected = next((item for item in plan.versions if item.id == version_id), None)
    if selected is None:
        raise PlatformValidationError(
            f"marketing plan version does not belong to this plan: {version_id}"
        )
    return selected


def _labels(locale: str) -> dict[str, str]:
    if locale == "en":
        return {
            "brand": "Heyu AI",
            "subtitle": "AI agricultural content marketing plan",
            "overview": "Plan overview",
            "product": "Product",
            "origin": "Origin",
            "platform": "Publishing platform",
            "audience": "Target audience",
            "version": "Version",
            "positioning": "Product positioning",
            "selling_points": "Core selling points",
            "strategy": "Platform strategy",
            "trend": "Trend integration",
            "topics": "Topic opportunities",
            "routes": "Creative routes and short-video scripts",
            "script": "Script",
            "music": "Background music",
            "shots": "Shot list",
            "cta": "Call to action",
            "quality": "Quality score",
            "livestream": "Livestream talking points",
            "calendar": "Seven-day content calendar",
            "actions": "Recommended next actions",
            "day": "Day {day}",
            "objective": "Objective",
            "content": "Content",
            "action": "Action",
        }
    if locale == "zh-HK":
        return {
            "brand": "禾語 AI",
            "subtitle": "AI 農產品內容營銷方案",
            "overview": "方案概覽",
            "product": "產品",
            "origin": "產地",
            "platform": "發佈平台",
            "audience": "目標受眾",
            "version": "版本",
            "positioning": "產品定位",
            "selling_points": "核心賣點",
            "strategy": "平台策略",
            "trend": "熱點融入",
            "topics": "選題機會",
            "routes": "創意路線與短影片腳本",
            "script": "文案腳本",
            "music": "背景音樂",
            "shots": "分鏡清單",
            "cta": "行動引導",
            "quality": "內容質素評分",
            "livestream": "直播話術",
            "calendar": "七日內容日曆",
            "actions": "下一步建議",
            "day": "第 {day} 日",
            "objective": "目標",
            "content": "內容",
            "action": "行動",
        }
    return {
        "brand": "禾语 AI",
        "subtitle": "AI 农产品内容营销方案",
        "overview": "方案概览",
        "product": "产品",
        "origin": "产地",
        "platform": "发布平台",
        "audience": "目标受众",
        "version": "版本",
        "positioning": "产品定位",
        "selling_points": "核心卖点",
        "strategy": "平台策略",
        "trend": "热点融入",
        "topics": "选题机会",
        "routes": "创意路线与短视频脚本",
        "script": "文案脚本",
        "music": "背景音乐",
        "shots": "分镜清单",
        "cta": "行动引导",
        "quality": "内容质量评分",
        "livestream": "直播话术",
        "calendar": "七天内容日历",
        "actions": "下一步建议",
        "day": "第 {day} 天",
        "objective": "目标",
        "content": "内容",
        "action": "行动",
    }


def _build_docx(
    plan: MarketingPlanDetailRead,
    version: MarketingPlanVersionRead,
    request: MarketingPlanRequest,
    content: MarketingPlanResponse,
) -> bytes:
    labels = _labels(request.locale)
    document = Document()
    document.styles["Normal"].font.name = "Arial"
    document.styles["Normal"].font.size = Pt(10.5)
    document.add_heading(labels["brand"], level=0)
    document.add_paragraph(labels["subtitle"])
    document.add_paragraph(plan.title)

    document.add_heading(labels["overview"], level=1)
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for key, value in (
        (labels["product"], request.product_name),
        (labels["origin"], request.origin or "-"),
        (labels["platform"], content.strategy.platform_name),
        (labels["audience"], request.audience or content.product_profile.core_audience),
        (labels["version"], f"v{version.version_number}"),
    ):
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value

    _docx_heading_with_text(
        document,
        labels["positioning"],
        content.product_profile.one_line_value,
    )
    document.add_paragraph(content.product_profile.story_angle)
    _docx_bullets(document, labels["selling_points"], content.product_profile.core_selling_points)

    document.add_heading(labels["strategy"], level=1)
    document.add_paragraph(content.strategy.content_focus)
    document.add_paragraph(
        f"{content.strategy.recommended_duration} · {content.strategy.conversion_action}"
    )

    document.add_heading(labels["trend"], level=1)
    document.add_paragraph(content.trend.trend_used)
    document.add_paragraph(content.trend.integration_method)
    document.add_paragraph(content.trend.caution)

    document.add_heading(labels["topics"], level=1)
    for topic in content.topic_signals:
        document.add_heading(f"{topic.title} · {topic.total_score}/100", level=2)
        document.add_paragraph(topic.content_angle)
        document.add_paragraph(topic.explanation)

    if content.videos:
        document.add_page_break()
        document.add_heading(labels["routes"], level=1)
        for index, video in enumerate(content.videos, start=1):
            document.add_heading(f"{index}. {video.title}", level=2)
            document.add_paragraph(video.hook)
            _docx_label_value(document, labels["script"], video.script)
            _docx_label_value(document, labels["music"], video.background_music)
            document.add_heading(labels["shots"], level=3)
            for shot in video.shots:
                document.add_paragraph(
                    f"{shot.seconds}｜{shot.visual}｜{shot.voiceover}｜{shot.filming_tip}",
                    style="List Number",
                )
            _docx_label_value(document, labels["cta"], video.call_to_action)
            _docx_label_value(
                document,
                labels["quality"],
                f"{video.quality_assessment.total_score}/100",
            )

    if content.livestream:
        document.add_heading(labels["livestream"], level=1)
        for section in content.livestream:
            document.add_heading(section.section, level=2)
            _docx_bullets(document, "", section.talking_points)

    if content.seven_day_plan:
        document.add_heading(labels["calendar"], level=1)
        calendar = document.add_table(rows=1, cols=4)
        calendar.style = "Table Grid"
        headers = calendar.rows[0].cells
        headers[0].text = labels["day"].replace(" {day}", "").replace("{day} ", "")
        headers[1].text = labels["objective"]
        headers[2].text = labels["content"]
        headers[3].text = labels["action"]
        for item in content.seven_day_plan:
            cells = calendar.add_row().cells
            cells[0].text = str(item.day)
            cells[1].text = item.objective
            cells[2].text = item.content
            cells[3].text = item.action

    _docx_bullets(document, labels["actions"], content.next_actions)
    payload = BytesIO()
    document.save(payload)
    return payload.getvalue()


def _docx_heading_with_text(document: DocxDocument, heading: str, text: str) -> None:
    document.add_heading(heading, level=1)
    document.add_paragraph(text)


def _docx_bullets(document: DocxDocument, heading: str, values: list[str]) -> None:
    if heading:
        document.add_heading(heading, level=2)
    for value in values:
        document.add_paragraph(value, style="List Bullet")


def _docx_label_value(document: DocxDocument, label: str, value: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.add_run(f"{label}：").bold = True
    paragraph.add_run(value)


def _build_pdf(
    plan: MarketingPlanDetailRead,
    version: MarketingPlanVersionRead,
    request: MarketingPlanRequest,
    content: MarketingPlanResponse,
) -> bytes:
    labels = _labels(request.locale)
    font_name = _pdf_font(request.locale)
    payload = BytesIO()
    document = SimpleDocTemplate(
        payload,
        pagesize=A4,
        rightMargin=18 * mm,
        leftMargin=18 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
        title=plan.title,
        author="Heyu AI",
    )
    base = getSampleStyleSheet()
    title = ParagraphStyle(
        "HeyuTitle",
        parent=base["Title"],
        fontName=font_name,
        fontSize=23,
        leading=31,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#17324d"),
        spaceAfter=8,
    )
    subtitle = ParagraphStyle(
        "HeyuSubtitle",
        parent=base["Normal"],
        fontName=font_name,
        fontSize=10,
        leading=16,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#4d6678"),
        spaceAfter=16,
    )
    heading = ParagraphStyle(
        "HeyuHeading",
        parent=base["Heading1"],
        fontName=font_name,
        fontSize=15,
        leading=22,
        textColor=colors.HexColor("#17324d"),
        spaceBefore=12,
        spaceAfter=7,
    )
    subheading = ParagraphStyle(
        "HeyuSubheading",
        parent=base["Heading2"],
        fontName=font_name,
        fontSize=11.5,
        leading=18,
        textColor=colors.HexColor("#167a66"),
        spaceBefore=8,
        spaceAfter=4,
    )
    body = ParagraphStyle(
        "HeyuBody",
        parent=base["BodyText"],
        fontName=font_name,
        fontSize=9.5,
        leading=16,
        textColor=colors.HexColor("#20313d"),
        spaceAfter=5,
    )
    story: list[Flowable] = [
        Paragraph(_xml(labels["brand"]), title),
        Paragraph(_xml(labels["subtitle"]), subtitle),
        Paragraph(_xml(plan.title), subtitle),
        Paragraph(_xml(labels["overview"]), heading),
    ]
    overview = [
        [labels["product"], request.product_name],
        [labels["origin"], request.origin or "-"],
        [labels["platform"], content.strategy.platform_name],
        [labels["audience"], request.audience or content.product_profile.core_audience],
        [labels["version"], f"v{version.version_number}"],
    ]
    table = Table(
        [[Paragraph(_xml(str(cell)), body) for cell in row] for row in overview],
        colWidths=[38 * mm, 118 * mm],
    )
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#eaf2f6")),
                ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#b8cbd5")),
                ("INNERGRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#d4e0e6")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 7),
                ("RIGHTPADDING", (0, 0), (-1, -1), 7),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    story.extend(
        [
            table,
            Paragraph(_xml(labels["positioning"]), heading),
            Paragraph(_xml(content.product_profile.one_line_value), subheading),
            Paragraph(_xml(content.product_profile.story_angle), body),
            Paragraph(_xml(labels["selling_points"]), subheading),
            _pdf_list(content.product_profile.core_selling_points, body),
            Paragraph(_xml(labels["strategy"]), heading),
            Paragraph(_xml(content.strategy.content_focus), body),
            Paragraph(
                _xml(
                    f"{content.strategy.recommended_duration} · "
                    f"{content.strategy.conversion_action}"
                ),
                body,
            ),
            Paragraph(_xml(labels["trend"]), heading),
            Paragraph(_xml(content.trend.trend_used), subheading),
            Paragraph(_xml(content.trend.integration_method), body),
            Paragraph(_xml(content.trend.caution), body),
            Paragraph(_xml(labels["topics"]), heading),
        ]
    )
    for topic in content.topic_signals:
        story.append(
            KeepTogether(
                [
                    Paragraph(_xml(f"{topic.title} · {topic.total_score}/100"), subheading),
                    Paragraph(_xml(topic.content_angle), body),
                    Paragraph(_xml(topic.explanation), body),
                ]
            )
        )
    if content.videos:
        story.extend([PageBreak(), Paragraph(_xml(labels["routes"]), heading)])
        for index, video in enumerate(content.videos, start=1):
            parts: list[Flowable] = [
                Paragraph(_xml(f"{index}. {video.title}"), subheading),
                Paragraph(_xml(video.hook), body),
                Paragraph(_xml(f"{labels['script']}：{video.script}"), body),
                Paragraph(_xml(f"{labels['music']}：{video.background_music}"), body),
                Paragraph(_xml(labels["shots"]), subheading),
            ]
            parts.extend(
                Paragraph(
                    _xml(f"{shot.seconds}｜{shot.visual}｜{shot.voiceover}｜{shot.filming_tip}"),
                    body,
                )
                for shot in video.shots
            )
            parts.extend(
                [
                    Paragraph(_xml(f"{labels['cta']}：{video.call_to_action}"), body),
                    Paragraph(
                        _xml(f"{labels['quality']}：{video.quality_assessment.total_score}/100"),
                        body,
                    ),
                    Spacer(1, 5 * mm),
                ]
            )
            story.append(KeepTogether(parts))
    if content.livestream:
        story.append(Paragraph(_xml(labels["livestream"]), heading))
        for section in content.livestream:
            story.extend(
                [
                    Paragraph(_xml(section.section), subheading),
                    _pdf_list(section.talking_points, body),
                ]
            )
    if content.seven_day_plan:
        story.append(Paragraph(_xml(labels["calendar"]), heading))
        for day in content.seven_day_plan:
            story.append(
                KeepTogether(
                    [
                        Paragraph(
                            _xml(f"{labels['day'].format(day=day.day)} · {day.objective}"),
                            subheading,
                        ),
                        Paragraph(_xml(day.content), body),
                        Paragraph(_xml(day.action), body),
                    ]
                )
            )
    story.extend(
        [
            Paragraph(_xml(labels["actions"]), heading),
            _pdf_list(content.next_actions, body),
        ]
    )
    document.build(story)
    return payload.getvalue()


def _pdf_font(locale: str) -> str:
    font_name = "MSung-Light" if locale == "zh-HK" else "STSong-Light"
    if font_name not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(UnicodeCIDFont(font_name))
    return font_name


def _pdf_list(values: list[str], style: ParagraphStyle) -> ListFlowable:
    items: list[Flowable] = [
        cast(Flowable, ListItem(Paragraph(_xml(value), style), leftIndent=8)) for value in values
    ]
    return ListFlowable(
        items,
        bulletType="bullet",
        leftIndent=16,
        bulletFontName=style.fontName,
    )


def _xml(value: str) -> str:
    return (
        value.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace("\n", "<br/>")
    )
