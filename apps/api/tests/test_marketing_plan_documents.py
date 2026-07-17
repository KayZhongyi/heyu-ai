from io import BytesIO

import pytest
from docx import Document

from tests.conftest import bootstrap
from tests.test_marketing_plan_library import plan_payload


@pytest.mark.parametrize(
    ("locale", "product_name"),
    [
        ("zh-CN", "当季番茄"),
        ("zh-HK", "當季番茄"),
        ("en", "Seasonal Tomatoes"),
    ],
)
def test_saved_plan_downloads_word_document(client, auth, locale, product_name):
    created = client.post(
        "/v1/marketing-plans",
        headers=auth,
        json=plan_payload(client, locale=locale, product_name=product_name),
    )
    assert created.status_code == 201, created.text

    response = client.get(
        f"/v1/marketing-plans/{created.json()['id']}/document",
        headers=auth,
        params={"format": "docx"},
    )

    assert response.status_code == 200, response.text
    assert response.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert response.headers["x-heyu-content-sha256"]
    assert 'filename="heyu-marketing-plan-v1.docx"' in response.headers["content-disposition"]
    document = Document(BytesIO(response.content))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert created.json()["title"] in text
    assert product_name in "\n".join(
        cell.text for table in document.tables for row in table.rows for cell in row.cells
    )


@pytest.mark.parametrize("locale", ["zh-CN", "zh-HK", "en"])
def test_saved_plan_downloads_pdf_document(client, auth, locale):
    created = client.post(
        "/v1/marketing-plans",
        headers=auth,
        json=plan_payload(client, locale=locale),
    )
    assert created.status_code == 201, created.text

    response = client.get(
        f"/v1/marketing-plans/{created.json()['id']}/document",
        headers=auth,
        params={"format": "pdf"},
    )

    assert response.status_code == 200, response.text
    assert response.headers["content-type"] == "application/pdf"
    assert response.content.startswith(b"%PDF-")
    assert response.headers["x-heyu-content-sha256"]
    assert 'filename="heyu-marketing-plan-v1.pdf"' in response.headers["content-disposition"]


def test_document_export_can_select_a_historical_version(client, auth):
    created = client.post(
        "/v1/marketing-plans",
        headers=auth,
        json=plan_payload(client),
    ).json()
    first_version = created["current_version"]
    updated_content = dict(first_version["content"])
    updated_content["product_profile"] = dict(updated_content["product_profile"])
    updated_content["product_profile"]["one_line_value"] = "Updated positioning"
    updated = client.post(
        f"/v1/marketing-plans/{created['id']}/versions",
        headers=auth,
        json={
            "request_payload": first_version["request_payload"],
            "content": updated_content,
            "change_summary": "Updated positioning",
        },
    )
    assert updated.status_code == 201, updated.text

    response = client.get(
        f"/v1/marketing-plans/{created['id']}/document",
        headers=auth,
        params={"format": "docx", "version_id": first_version["id"]},
    )

    assert response.status_code == 200, response.text
    assert 'filename="heyu-marketing-plan-v1.docx"' in response.headers["content-disposition"]


def test_document_export_rejects_invalid_format_and_foreign_version(client, auth):
    first = client.post(
        "/v1/marketing-plans",
        headers=auth,
        json=plan_payload(client, product_name="First Product"),
    ).json()
    second = client.post(
        "/v1/marketing-plans",
        headers=auth,
        json=plan_payload(client, product_name="Second Product"),
    ).json()

    invalid_format = client.get(
        f"/v1/marketing-plans/{first['id']}/document",
        headers=auth,
        params={"format": "txt"},
    )
    foreign_version = client.get(
        f"/v1/marketing-plans/{first['id']}/document",
        headers=auth,
        params={"format": "pdf", "version_id": second["current_version"]["id"]},
    )

    assert invalid_format.status_code == 422
    assert foreign_version.status_code == 422
    assert "does not belong to this plan" in foreign_version.json()["detail"]


def test_document_export_requires_authentication_and_preserves_tenant_isolation(
    client,
    auth,
):
    created = client.post(
        "/v1/marketing-plans",
        headers=auth,
        json=plan_payload(client),
    ).json()
    second_owner = bootstrap(client, "document-export-other", "other@example.com")
    second_auth = {"Authorization": f"Bearer {second_owner['access_token']}"}

    unauthenticated = client.get(
        f"/v1/marketing-plans/{created['id']}/document",
        params={"format": "pdf"},
    )
    foreign_tenant = client.get(
        f"/v1/marketing-plans/{created['id']}/document",
        headers=second_auth,
        params={"format": "pdf"},
    )

    assert unauthenticated.status_code == 401
    assert foreign_tenant.status_code == 404
